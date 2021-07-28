from docx import Document
import cx_Oracle
import os
import json

#SETUP TEST(S) HERE
params = ['SARSH']
#END SETUP TEST(S)

#SETUP TEMPLATE HERE
template = r'dictionaries\templates\antigen.docx'
#END SETUP TEMPLATE

#convert test(s) array to string
param_in_string = ','.join("'{}'".format(param) for param in params)


def getdata(filename, labno):
    #get database config
    setting_file = os.path.abspath(os.path.join(os.path.dirname(__file__),"..","setting.json"))
    with open(setting_file,'r') as f :
        config = json.load(f)
    user_db = config["database"]["user"]
    pass_db = config["database"]["pass"]
    host_db = config["database"]["host"]

    try:
        conn = cx_Oracle.connect(user_db,pass_db,host_db)
        cursor = conn.cursor()
        sql = """
            select 
            to_char(oh_trx_dt,'dd-MON-yyyy'), oh_tno, oh_apid, oh_last_name, 
            (case when oh_pataddr1 is null then '-' else oh_pataddr1 end), 
            (case when oh_pataddr3 is null then '-' else oh_pataddr3 end), 
            (case when oh_pataddr4 is null then '-' else oh_pataddr4 end), 
            to_char(oh_bod,'dd-MON-yyyy'), 
        """
        for param in params:
            sql = sql + "max(decode(od_testcode,'{}',(case tv_desc when null then od_tr_val else tv_desc end))){},".format(param,param)
        sql = sql + f"""
            (case oh_sex when '1' then 'Laki-laki' when '2' then 'Perempuan' end)oh_sex
            from hord_hdr
            inner join hord_dtl on oh_tno = od_tno
            left join textvalue on tv_code = od_tr_val
            where oh_tno = '{labno}'
            and od_testcode in({param_in_string})
            group by oh_trx_dt, oh_tno, oh_apid, oh_last_name, oh_pataddr1, oh_pataddr3, oh_pataddr4,
            oh_bod, oh_sex
        """

        print(sql)
        
        cursor.execute(sql)
        record = cursor.fetchone()

        #SETUP VARIABLE VALUE HERE
        replacements = {
            '{trx_dt}' : record[0],
            '{apid}': record[2],
            '{nama}': record[3],
            '{address1}' : record[4],
            '{address3}' : record[5],
            '{address4}' : record[6],
            '{dob}': record[7],
            '{value}': record[8],
            '{sex}': record[9],
            }
        #END SETUP VARIABLE VALUE
        
        #USE THIS LOOP IF DOCX TEMPLATE USING PARAGRAPH FORMAT
        #get word data and replace specific word based on replacement variable
        doc = Document(template)
        #for paragraph in doc.paragraphs:
        #    for key in replacements:
        #        paragraph.text = paragraph.text.replace(key, replacements[key])
        ######################################################################################
        
        
        #USE THIS LOOP IF DOCX TEMPLATE USING TABLE FORMAT
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for key in replacements:
                            if key in paragraph.text:
                                paragraph.text = paragraph.text.replace(key,replacements[key])
        #########################################################################################
        doc.save(filename)
        
    except cx_Oracle.Error as e:
        print(e)

if __name__ == "__main__":
    getdata("D:\\test.docx","21000001")